"""Reader for Word ``.docx`` files — extracts content as Markdown.

Uses ``python-docx`` to walk paragraphs and runs, preserving:

* Headings (Word styles ``Heading 1``–``Heading 6`` → Markdown ``#``–``######``)
* Bold, italic, inline code (if font is Consolas / Courier New)
* Ordered and unordered lists (Word list paragraphs)
* Tables (basic grid)
* Images (placeholder links)
* Paragraph breaks
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Optional

from .base import DocumentReader
from dolatex.format import DocumentFormat


class DocxReader(DocumentReader):
    """Convert a ``.docx`` file to Markdown text."""

    def extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    def read(self, path: Path | str | bytes) -> str:
        """Read a .docx file and return its content as Markdown."""
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ModuleNotFoundError:
            raise ValueError(
                "python-docx is required to read .docx files. "
                "Install with: pip install dolatex[docx] or pip install python-docx"
            )

        if isinstance(path, bytes):
            import io
            doc = DocxDocument(io.BytesIO(path))
        else:
            doc = DocxDocument(str(path))

        lines: list[str] = []
        list_counter: dict[int, int] = {}
        in_table = False
        table_capture: list[list[str]] = []

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            text = self._runs_to_markdown(para)

            # --- Inline images from runs --------------------------------------
            # python-docx doesn't expose images in runs easily via the public
            # API, so we walk the XML manually.
            img_md = self._extract_inline_images(para)
            if img_md and text:
                text += "\n" + img_md
            elif img_md:
                text = img_md

            # --- Headings -----------------------------------------------------
            if style_name.startswith("heading"):
                level = self._parse_heading_level(style_name, para)
                lines.append(f"{'#' * level} {text}")
                lines.append("")
                continue

            # --- Tables (captured separately, see table handling below) -------
            # Tables in docx are represented as XML inside paragraphs.
            table_xml = para._element.find(qn("w:r"))
            # We handle tables via doc.tables below — skip paragraphs that are
            # purely table cells.

            # --- Lists --------------------------------------------------------
            numPr = para._element.find(qn("w:pPr"))
            ilvl_elem = None
            num_id_elem = None
            if numPr is not None:
                numPr_inner = numPr.find(qn("w:numPr"))
                if numPr_inner is not None:
                    ilvl_elem = numPr_inner.find(qn("w:ilvl"))
                    num_id_elem = numPr_inner.find(qn("w:numId"))

            if num_id_elem is not None and ilvl_elem is not None:
                ilvl_val = ilvl_elem.get(qn("w:val")) if ilvl_elem is not None else "0"
                ilvl = int(ilvl_val) if ilvl_val else 0
                indent = "  " * ilvl

                # Determine ordered vs unordered
                num_id_val = num_id_elem.get(qn("w:val"))
                is_ordered = self._is_ordered_list(doc, num_id_val)

                if is_ordered:
                    list_counter[ilvl] = list_counter.get(ilvl, 0) + 1
                    lines.append(f"{indent}{list_counter[ilvl]}. {text}")
                    # Reset counters for deeper levels
                    for k in list(list_counter):
                        if k > ilvl:
                            del list_counter[k]
                else:
                    lines.append(f"{indent}- {text}")
                continue

            # --- Plain paragraph ----------------------------------------------
            if text.strip():
                lines.append(text)
                lines.append("")
            else:
                lines.append("")

        # --- Tables -----------------------------------------------------------
        for table in doc.tables:
            md_rows: list[str] = []
            for row_idx, row in enumerate(table.rows):
                cell_texts: list[str] = []
                for cell in row.cells:
                    cell_paras = [
                        self._runs_to_markdown(p)
                        for p in cell.paragraphs
                        if p.text.strip()
                    ]
                    cell_texts.append(" ".join(cell_paras) if cell_paras else cell.text)
                md_rows.append("| " + " | ".join(cell_texts) + " |")
                if row_idx == 0:
                    md_rows.append(
                        "| " + " | ".join("---" for _ in row.cells) + " |"
                    )
            if md_rows:
                lines.append("")
                lines.extend(md_rows)
                lines.append("")

        return "\n".join(lines).strip()

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _runs_to_markdown(para: Any) -> str:
        """Convert the runs of a paragraph to inline Markdown."""
        parts: list[str] = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            # Inline code heuristic: Consolas / Courier New font
            font_name = (run.font.name or "").lower()
            if "consolas" in font_name or "courier" in font_name:
                text = f"`{text}`"
            parts.append(text)
        return "".join(parts)

    @staticmethod
    def _extract_inline_images(para: Any) -> str:
        """Extract ``drawing`` XML children as Markdown image links."""
        from docx.oxml.ns import qn

        md_parts: list[str] = []
        for drawing in para._element.iter(qn("w:drawing")):
            blip = drawing.find(".//" + qn("a:blip"))
            if blip is not None:
                embed = blip.get(qn("r:embed"))
                if embed:
                    md_parts.append(f"![Image]({embed})")
        return "\n".join(md_parts) if md_parts else ""

    @staticmethod
    def _parse_heading_level(style_name: str, para: Any) -> int:
        """Map ``heading 1`` → 1, ``heading 2`` → 2, etc."""
        for i in range(1, 7):
            if str(i) in style_name:
                return i
        # Fallback: check outline level
        from docx.oxml.ns import qn

        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            outline = pPr.find(qn("w:outlineLvl"))
            if outline is not None:
                val = outline.get(qn("w:val"))
                if val:
                    return int(val) + 1
        return 1

    @staticmethod
    def _is_ordered_list(doc: Any, num_id_val: Optional[str]) -> bool:
        """Check whether a list ``numId`` is ordered (decimal) or bullet."""
        from docx.oxml.ns import qn

        if not num_id_val:
            return False
        # Walk numbering definitions in the XML
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        for num in numbering.findall(qn("w:num")):
            if num.get(qn("w:numId")) == num_id_val:
                ab = num.find(qn("w:abstractNumId"))
                if ab is not None:
                    abs_id = ab.get(qn("w:val"))
                    # Look up abstract numbering format
                    for abs_num in numbering.findall(qn("w:abstractNum")):
                        if abs_num.get(qn("w:abstractNumId")) == abs_id:
                            lvl = abs_num.find(qn("w:lvl"))
                            if lvl is not None:
                                numFmt = lvl.find(qn("w:numFmt"))
                                if numFmt is not None:
                                    fmt = numFmt.get(qn("w:val"), "")
                                    return fmt == "decimal"
        return False  # default to unordered

    def extract_format(self, data: str | bytes | Path) -> Optional[DocumentFormat]:
        """Extract page layout and typography info from a .docx file."""
        try:
            from docx import Document as DocxDocument
        except ModuleNotFoundError:
            return DocumentFormat()  # defaults

        if isinstance(data, bytes):
            import io

            doc = DocxDocument(io.BytesIO(data))
        elif isinstance(data, Path):
            doc = DocxDocument(str(data))
        else:
            doc = DocxDocument(str(data))

        fmt = DocumentFormat()

        # --- Page size & margins (from first section) ---
        section = doc.sections[0] if doc.sections else None
        if section:
            fmt.page_size = self._page_size_str(section.page_width, section.page_height)
            fmt.orientation = "landscape" if section.orientation else "portrait"
            fmt.margin_top = self._emu_to_cm_str(section.top_margin)
            fmt.margin_bottom = self._emu_to_cm_str(section.bottom_margin)
            fmt.margin_left = self._emu_to_cm_str(section.left_margin)
            fmt.margin_right = self._emu_to_cm_str(section.right_margin)

        # --- Default font (from style defaults) ---
        try:
            style = doc.styles["Normal"]
            if style.font and style.font.name:
                fmt.font_family = style.font.name.lower()
            if style.font and style.font.size:
                fmt.body_font_size = self._emu_to_pt_str(style.font.size)
        except (KeyError, AttributeError):
            pass

        # --- Default paragraph format ---
        try:
            style = doc.styles["Normal"]
            pf = style.paragraph_format
            if pf.alignment is not None:
                fmt.alignment = self._alignment_str(pf.alignment)
            if pf.first_line_indent:
                fmt.first_line_indent = self._emu_to_cm_str(pf.first_line_indent)
        except (KeyError, AttributeError):
            pass

        # --- Heading font sizes ---
        for level in range(1, 5):
            try:
                h_style = doc.styles[f"Heading {level}"]
                if h_style.font and h_style.font.size:
                    fmt.heading_sizes[level] = self._emu_to_pt_str(h_style.font.size)
            except (KeyError, AttributeError):
                pass

        return fmt

    @staticmethod
    def _emu_to_cm_str(emu: int) -> str:
        """Convert EMU (English Metric Units) to a cm string like "2.54cm"."""
        cm = emu / 360000  # 1 cm = 360000 EMU
        return f"{cm:.1f}cm"

    @staticmethod
    def _emu_to_pt_str(emu: int) -> str:
        """Convert EMU (English Metric Units) to a pt string like "12pt"."""
        pt = round(emu / 12700)  # 1 pt = 12700 EMU
        return f"{pt}pt"

    @staticmethod
    def _page_size_str(width_emu: int, height_emu: int) -> str:
        """Map DOCX page dimensions to LaTeX page size names."""
        # A4: 210x297mm = ~7560000x10692000 EMU
        # Letter: 215.9x279.4mm = ~7772400x10058400 EMU
        w_mm = width_emu / 36000  # 1 mm = 36000 EMU
        h_mm = height_emu / 36000
        if abs(w_mm - 210) < 5 and abs(h_mm - 297) < 5:
            return "a4paper"
        if abs(w_mm - 215.9) < 5 and abs(h_mm - 279.4) < 5:
            return "letterpaper"
        # Custom: use width x height
        return f"{w_mm:.0f}mm x {h_mm:.0f}mm"

    @staticmethod
    def _alignment_str(alignment: int) -> str:
        """Map python-docx WD_ALIGN_PARAGRAPH constants to strings."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justified",
        }
        return mapping.get(alignment, "justified")
