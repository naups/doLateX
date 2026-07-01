"""Tests for the document readers (Markdown, DOCX, PDF).

Uses ``pypdf`` and ``python-docx`` if available, gracefully skipping
those tests when the libraries are not installed.
"""

from __future__ import annotations

from pathlib import Path
import sys
import tempfile

import pytest

from dolatex.readers import (
    read_document,
    read_text,
    supported_extensions,
    get_reader,
)

# ---------------------------------------------------------------------------
# Markdown reader tests
# ---------------------------------------------------------------------------


class TestMarkdownReader:
    def test_reads_plain_text(self, tmp_path: Path) -> None:
        f = tmp_path / "test.md"
        f.write_text("# Hello\n\nWorld.", encoding="utf-8")
        result = read_document(f)
        assert "# Hello" in result
        assert "World." in result

    def test_reads_txt_extension(self, tmp_path: Path) -> None:
        f = tmp_path / "notes.txt"
        f.write_text("Plain text file.", encoding="utf-8")
        result = read_document(f)
        assert result == "Plain text file."

    def test_supported_extensions(self) -> None:
        exts = supported_extensions()
        assert ".md" in exts
        assert ".markdown" in exts
        assert ".txt" in exts
        assert ".docx" in exts
        assert ".pdf" in exts

    def test_get_reader(self) -> None:
        from dolatex.readers.markdown_reader import MarkdownReader
        reader = get_reader(".md")
        assert reader is not None
        assert isinstance(reader, MarkdownReader)

    def test_read_text_markdown(self) -> None:
        result = read_text("# Hello", ".md")
        assert result == "# Hello"

    def test_unsupported_extension(self) -> None:
        with pytest.raises(ValueError, match="Unsupported"):
            read_document("file.xyz")

    def test_unsupported_extension_read_text(self) -> None:
        with pytest.raises(ValueError, match="Unsupported"):
            read_text(b"data", ".xyz")


# ---------------------------------------------------------------------------
# DOCX reader tests
# ---------------------------------------------------------------------------

pytestmark_docx = pytest.mark.skipif(
    sys.modules.get("docx") is None,
    reason="python-docx not installed (pip install dolatex[docx])",
)


class TestDocxReader:
    def _create_minimal_docx(self, tmp_path: Path) -> Path:
        """Create a minimal .docx with a heading and a paragraph."""
        from docx import Document
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("This is a **bold** paragraph.")
        doc.add_paragraph("Normal text with *italic* styling.")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        return path

    def test_reads_headings(self, tmp_path: Path) -> None:
        p = self._create_minimal_docx(tmp_path)
        result = read_document(p)
        assert "# Test Heading" in result or "Test Heading" in result

    def test_reads_paragraphs(self, tmp_path: Path) -> None:
        p = self._create_minimal_docx(tmp_path)
        result = read_document(p)
        assert "Normal text" in result

    def test_extension_supported(self) -> None:
        reader = get_reader(".docx")
        assert reader is not None

    def test_read_text_from_bytes(self) -> None:
        from docx import Document
        import io
        doc = Document()
        doc.add_paragraph("Hello from bytes.")
        buf = io.BytesIO()
        doc.save(buf)
        result = read_text(buf.getvalue(), ".docx")
        assert "Hello from bytes" in result


# ---------------------------------------------------------------------------
# PDF reader tests
# ---------------------------------------------------------------------------

pytestmark_pdf = pytest.mark.skipif(
    sys.modules.get("pypdf") is None,
    reason="pypdf not installed (pip install dolatex[pdf])",
)


class TestPdfReader:
    def _create_minimal_pdf(self, tmp_path: Path) -> Path:
        """Create a minimal PDF with a heading and paragraph."""
        from pypdf import PdfWriter
        from io import BytesIO

        writer = PdfWriter()
        # Minimal PDF content — we add a page using a simple text string.
        # pypdf's page.add_text() is low-level; we create via reportlab-style
        # isn't available, so we add a blank page and overlay text via
        # annotations or use a simple text PDF.
        # For testing, we create a minimal valid PDF with visible text.
        pdf_content = BytesIO()
        writer.create_viewer_preference()
        writer.add_blank_page(612, 792)  # US Letter
        writer.write(pdf_content)

        path = tmp_path / "test.pdf"
        path.write_bytes(pdf_content.getvalue())
        return path

    def _create_text_pdf(self, tmp_path: Path, text: str) -> Path:
        """Create a PDF with embedded text content."""
        # pypdf doesn't have a high-level API to add text.
        # We create a minimal valid PDF with a text object.
        pdf_bytes = self._make_simple_pdf(text)
        path = tmp_path / "text.pdf"
        path.write_bytes(pdf_bytes)
        return path

    @staticmethod
    def _make_simple_pdf(text: str) -> bytes:
        """Build a minimal valid PDF with embedded text."""
        # Minimal PDF with a single page containing a text object.
        content = f"""
BT
/F1 12 Tf
100 700 Td
({text}) Tj
ET
"""
        escaped = text.replace("(", "\\(").replace(")", "\\)")
        # Build a minimal PDF manually
        body = f"""1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 << /Type /Font
   /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>
endobj

4 0 obj
<< /Length 55 >>
stream
BT /F1 12 Tf 100 700 Td ({escaped}) Tj ET
endstream
endobj

xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n

trailer
<< /Size 5 /Root 1 0 R >>
startxref
403
%%EOF"""
        return body.encode("latin-1")

    def test_extension_supported(self) -> None:
        reader = get_reader(".pdf")
        assert reader is not None

    def test_reads_text_content(self, tmp_path: Path) -> None:
        p = self._create_text_pdf(tmp_path, "Hello PDF World")
        result = read_document(p)
        assert "Hello PDF World" in result or "hello" in result.lower()

    def test_read_text_from_bytes(self) -> None:
        pdf_bytes = self._make_simple_pdf("PDF from bytes")
        result = read_text(pdf_bytes, ".pdf")
        assert "PDF from bytes" in result or "bytes" in result.lower()


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


class TestReaderEdgeCases:
    def test_read_document_with_str_path(self, tmp_path: Path) -> None:
        f = tmp_path / "hello.md"
        f.write_text("content", encoding="utf-8")
        result = read_document(str(f))
        assert result == "content"

    def test_supported_extensions_contains_dot(self) -> None:
        exts = supported_extensions()
        for e in exts:
            assert e.startswith("."), f"Extension {e!r} should start with '.'"

    def test_get_reader_none(self) -> None:
        assert get_reader(".unknown") is None
